"""Test helpers for constructing `.docx` fixtures with comments."""

from __future__ import annotations

import typing as typ

from docx import Document

if typ.TYPE_CHECKING:
    from pathlib import Path

    from docx.comments import Comment as WordComment


class CommentElement(typ.Protocol):
    """The subset of the private comment element needed by the fixtures."""

    def set(self, key: str, value: str) -> None:
        """Assign an XML attribute value."""


COMMENT_DATE_ATTR = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date"
FIXED_COMMENT_TIMESTAMP = "2026-04-09T20:35:31Z"
NAIVE_COMMENT_TIMESTAMP = "2026-04-09T20:35:31"


def build_simple_comment_docx(path: Path) -> Path:
    """Create a document with a heading and one commented run."""
    document = Document()
    document.add_heading("Sample heading", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Before ")
    commented_run = paragraph.add_run("commented text")
    paragraph.add_run(" after.")
    comment = document.add_comment(
        commented_run,
        text="Needs evidence.",
        author="Sam C",
    )
    _set_fixed_timestamp(comment)
    document.save(str(path))
    return path


def build_multi_run_comment_docx(path: Path) -> Path:
    """Create a document whose comment spans multiple runs."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Please review ")
    first_run = paragraph.add_run("this ")
    middle_run = paragraph.add_run("multi-run ")
    last_run = paragraph.add_run("span")
    paragraph.add_run(" now.")
    comment = document.add_comment(
        [first_run, middle_run, last_run],
        text="Tighten this wording.",
        author="Alex Reviewer",
    )
    _set_fixed_timestamp(comment)
    document.save(str(path))
    return path


def build_cross_paragraph_comment_docx(path: Path) -> Path:
    """Create a document whose comment spans two paragraphs."""
    document = Document()
    first_paragraph = document.add_paragraph()
    first_paragraph.add_run("The first paragraph opens ")
    first_anchor = first_paragraph.add_run("a long comment")
    second_paragraph = document.add_paragraph()
    second_anchor = second_paragraph.add_run("that closes here")
    second_paragraph.add_run(" with more context.")
    comment = document.add_comment(
        [first_anchor, second_anchor],
        text="This crosses a paragraph boundary.",
        author="Jordan",
    )
    _set_fixed_timestamp(comment)
    comment.add_paragraph("Second note paragraph.")
    document.save(str(path))
    return path


def build_table_docx(path: Path) -> Path:
    """Create a document with an unsupported table block."""
    document = Document()
    document.add_paragraph("Before table.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Ignored"
    document.add_paragraph("After table.")
    document.save(str(path))
    return path


def build_criticmarkup_literal_docx(path: Path) -> Path:
    """Create a document containing literal CriticMarkup delimiters."""
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Literal {==marker==} text")
    comment = document.add_comment(
        run,
        text="Avoid {>>raw<<} delimiters.",
        author="Escaper",
    )
    _set_fixed_timestamp(comment)
    document.save(str(path))
    return path


def build_comment_normalization_docx(path: Path) -> Path:
    """Create a comment with blank authorship and a naive timestamp."""
    document = Document()
    run = document.add_paragraph().add_run("Commented text")
    comment = document.add_comment(run, text="Normalize metadata.", author="   ")
    _set_timestamp(comment, NAIVE_COMMENT_TIMESTAMP)
    document.save(str(path))
    return path


def build_fixture(name: str, path: Path) -> Path:
    """Build a named fixture document at `path`."""
    builders = {
        "simple-comment": build_simple_comment_docx,
        "multi-run-comment": build_multi_run_comment_docx,
        "cross-paragraph-comment": build_cross_paragraph_comment_docx,
        "table-document": build_table_docx,
        "criticmarkup-literal": build_criticmarkup_literal_docx,
        "comment-normalization": build_comment_normalization_docx,
    }
    return builders[name](path)


def _set_fixed_timestamp(comment: WordComment) -> None:
    """Force a deterministic comment timestamp for snapshot stability."""
    _set_timestamp(comment, FIXED_COMMENT_TIMESTAMP)


def _set_timestamp(comment: WordComment, value: str) -> None:
    """Set the private OOXML timestamp used by a synthetic fixture."""
    element = typ.cast(
        "CommentElement",
        object.__getattribute__(comment, "_element"),
    )
    element.set(COMMENT_DATE_ATTR, value)
