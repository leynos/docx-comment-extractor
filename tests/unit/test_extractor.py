"""Unit tests for the document extractor."""

from __future__ import annotations

import datetime as dt
import typing as typ
import zlib

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
from lxml.etree import XMLSyntaxError

from docx_comment_extractor import extractor
from docx_comment_extractor.extractor import (
    ExtractionError,
    _extract_paragraph_block,
    extract_document,
)
from docx_comment_extractor.renderer import render_document
from tests.support_documents import build_fixture

if typ.TYPE_CHECKING:
    from pathlib import Path

    from docx.document import Document as WordDocument


def test_extract_document_uses_injected_document_loader(tmp_path: Path) -> None:
    """The public query should delegate package I/O to its loader boundary."""
    document_path = build_fixture("simple-comment", tmp_path / "simple.docx")
    loaded_paths: list[Path] = []

    def load_document(path: Path) -> WordDocument:
        """Record and load the requested package."""
        loaded_paths.append(path)
        return Document(str(path))

    result = extract_document(document_path, document_loader=load_document)

    assert loaded_paths == [document_path], (
        "the injected loader should receive the path"
    )
    assert result.document.comments, "extraction should use the injected document"


def test_extract_document_wraps_loader_failures(tmp_path: Path) -> None:
    """Loader failures should cross the public API as an extraction error."""
    document_path = tmp_path / "broken.docx"
    document_path.touch()

    def fail_to_load(path: Path) -> typ.NoReturn:
        """Simulate an unreadable document package."""
        del path
        message = "storage detail"
        raise OSError(message)

    with pytest.raises(ExtractionError, match="Could not extract the Word document"):
        extract_document(document_path, document_loader=fail_to_load)


def _create_xml_syntax_error() -> XMLSyntaxError:
    """Produce the parser-originated XML exception used by a corrupt package."""
    try:
        etree.fromstring(b"<")
    except XMLSyntaxError as error:
        return error
    message = "The malformed XML fixture should raise XMLSyntaxError."
    raise AssertionError(message)


@pytest.mark.parametrize(
    "failure_factory",
    [
        _create_xml_syntax_error,
        lambda: zlib.error("invalid compressed payload"),
        lambda: EOFError("unexpected end of package"),
        lambda: RuntimeError("unexpected package state"),
    ],
)
def test_extract_document_wraps_corrupt_package_failures(
    tmp_path: Path,
    failure_factory: typ.Callable[[], Exception],
) -> None:
    """Known corrupt-package errors should not cross the public API boundary."""
    document_path = tmp_path / "corrupt.docx"
    document_path.touch()

    def fail_to_load(path: Path) -> typ.NoReturn:
        del path
        raise failure_factory()

    with pytest.raises(ExtractionError, match="Could not extract the Word document"):
        extract_document(document_path, document_loader=fail_to_load)


def test_extract_document_rejects_oversized_packages(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """The public API should reject oversized packages before loading them."""
    document_path = tmp_path / "oversized.docx"
    document_path.write_bytes(b"xx")
    monkeypatch.setattr(extractor, "MAX_INPUT_BYTES", 1)

    def fail_if_loaded(path: Path) -> typ.NoReturn:
        del path
        pytest.fail("oversized packages should be rejected before loader invocation")

    with pytest.raises(ExtractionError, match="too large"):
        extract_document(document_path, document_loader=fail_if_loaded)


def test_extract_document_builds_simple_model(tmp_path: Path) -> None:
    """A single commented run should be represented in the block model."""
    document_path = build_fixture("simple-comment", tmp_path / "simple.docx")

    result = extract_document(document_path)

    assert result.warnings == (), "the simple model should not contain warnings"
    assert [block.kind for block in result.document.blocks] == [
        "heading",
        "paragraph",
    ], "the simple model should preserve heading and paragraph block kinds"
    assert result.document.comments[0].author == "Sam C", (
        "the extracted comment should preserve its normalized author"
    )
    assert result.document.comments[0].body == "Needs evidence.", (
        "the extracted comment should preserve its body"
    )
    paragraph_fragments = result.document.blocks[1].fragments
    assert paragraph_fragments[1].start_comment_ids == ("0",), (
        "the commented fragment should open comment boundary 0"
    )
    assert paragraph_fragments[1].end_comment_ids == ("0",), (
        "the commented fragment should close comment boundary 0"
    )


def test_extract_document_normalizes_comment_metadata(tmp_path: Path) -> None:
    """Blank authors and naive timestamps should normalize predictably."""
    document_path = build_fixture(
        "comment-normalization",
        tmp_path / "comment-normalization.docx",
    )

    result = extract_document(document_path)

    assert result.warnings == (), "metadata normalization should not emit warnings"
    comment = result.document.comments[0]
    assert comment.author is None, "a whitespace-only author should normalize to None"
    assert comment.timestamp == dt.datetime(
        2026,
        4,
        9,
        20,
        35,
        31,
        tzinfo=dt.UTC,
    ), "a naive comment timestamp should normalize to UTC"


def test_extract_document_supports_multi_run_ranges(tmp_path: Path) -> None:
    """A comment spanning multiple runs should start and end on separate fragments."""
    document_path = build_fixture("multi-run-comment", tmp_path / "multi-run.docx")

    result = extract_document(document_path)

    paragraph_fragments = result.document.blocks[0].fragments
    assert paragraph_fragments[1].start_comment_ids == ("0",), (
        "the first commented run should open the multi-run boundary"
    )
    assert paragraph_fragments[3].end_comment_ids == ("0",), (
        "the last commented run should close the multi-run boundary"
    )


def test_extract_document_preserves_inline_controls_in_ranges(
    tmp_path: Path,
) -> None:
    """Tabs and line breaks should survive inside and outside comment ranges."""
    document_path = build_fixture("inline-controls", tmp_path / "controls.docx")

    result = extract_document(document_path)
    fragments = result.document.blocks[0].fragments

    assert [fragment.text for fragment in fragments] == [
        "Outside\t\n",
        "Inside\t\n",
        "again\n",
        "After\t\n",
    ], "the extracted model should preserve tabs and both line-break elements"
    assert fragments[1].start_comment_ids == ("0",), (
        "the first control-bearing fragment should open the comment range"
    )
    assert fragments[2].end_comment_ids == ("0",), (
        "the final control-bearing fragment should close the comment range"
    )
    assert "{==Inside\t\nagain\n==}" in render_document(result.document), (
        "the renderer should preserve controls inside the CriticMarkup highlight"
    )
    assert render_document(result.document).startswith("Outside\t\n"), (
        "the renderer should preserve controls outside the comment range"
    )


def test_extract_paragraph_block_accumulates_many_end_markers() -> None:
    """A fragment should collect many end markers without repeated tuple copying."""
    document = Document()
    paragraph = document.add_paragraph("marked")
    for marker_id in range(128):
        marker = OxmlElement("w:commentRangeEnd")
        marker.set(qn("w:id"), str(marker_id))
        # The fixture requires an OOXML-only comment boundary marker.
        paragraph._p.append(marker)

    block = _extract_paragraph_block(paragraph)

    assert block.fragments[0].end_comment_ids == tuple(map(str, range(128))), (
        "the paragraph block should retain every end marker in document order"
    )


def test_extract_document_supports_cross_paragraph_ranges(tmp_path: Path) -> None:
    """A comment spanning paragraphs should remain open until the second block."""
    document_path = build_fixture("cross-paragraph-comment", tmp_path / "cross.docx")

    result = extract_document(document_path)

    first_block = result.document.blocks[0]
    second_block = result.document.blocks[1]
    assert first_block.fragments[1].start_comment_ids == ("0",), (
        "the first block should open the cross-paragraph boundary"
    )
    assert second_block.fragments[0].end_comment_ids == ("0",), (
        "the second block should close the cross-paragraph boundary"
    )
    assert result.document.comments[0].body == (
        "This crosses a paragraph boundary. / Second note paragraph."
    ), "the cross-paragraph comment body should retain both normalized paragraphs"


def test_extract_document_warns_for_tables(tmp_path: Path) -> None:
    """Unsupported top-level tables should emit a warning rather than failing."""
    document_path = build_fixture("table-document", tmp_path / "table.docx")

    result = extract_document(document_path)

    assert len(result.warnings) == 1, "one table should produce one extraction warning"
    assert result.warnings[0].code == "unsupported-block", (
        "the table warning should use the unsupported-block code"
    )
    assert [
        "".join(fragment.text for fragment in block.fragments)
        for block in result.document.blocks
    ] == [
        "Before table.",
        "After table.",
    ], "table skipping should preserve the surrounding paragraph order"
